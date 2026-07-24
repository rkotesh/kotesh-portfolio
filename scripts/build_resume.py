from pathlib import Path
from textwrap import wrap

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer


ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "Sankula_Koteswara_Rao_ATS_Resume.docx"
PDF_PATH = ROOT / "assets" / "Sankula_Koteswara_Rao_Resume.pdf"


RESUME = {
    "name": "Sankula Koteswara Rao",
    "headline": "Software Developer | Python & Django Developer | Full-Stack Web Developer",
    "contact": (
        "Bapatla, Andhra Pradesh, India | +91 9182015717 | "
        "srkotesh23@gmail.com | linkedin.com/in/sankula-koteswararao | github.com/rkotesh | "
        "kotesh-portfolio-nine.vercel.app"
    ),
    "summary": (
        "Software developer and final-year B.Tech Artificial Intelligence and Machine Learning student "
        "with hands-on experience building Python, Django, Flask, REST API, React, JavaScript, and "
        "Streamlit applications. Completed a 6-month Python Development internship with Flipkart "
        "Launchpad and built public projects across employee leave management, ERP portals, AI chatbots, "
        "automation tools, and responsive web platforms. Strong fit for entry-level software developer, "
        "Python developer, backend developer, frontend developer, and full-stack developer roles."
    ),
    "skills": [
        "Languages: Python, JavaScript, TypeScript, HTML5, CSS3",
        "Backend: Django, Flask, REST APIs, authentication, API integration, server-side application logic",
        "Frontend: React.js, responsive UI, Bootstrap, Tailwind CSS, DOM manipulation, accessibility basics",
        "Databases and Tools: SQLite, Git, GitHub, Vercel, Render, Streamlit, VS Code",
        "AI and Automation: AI tools, prompt engineering workflows, chatbot development, Python automation",
        "Software Practices: debugging, clean code, version control, deployment, documentation, teamwork",
    ],
    "experience": [
        {
            "title": "Coordinator",
            "company": "Techno Future India",
            "date": "May 2026 - Present",
            "location": "Guntur, Andhra Pradesh",
            "bullets": [
                "Mentor internship cohorts on MERN Stack fundamentals, software development workflow, debugging, and project execution.",
                "Design structured learning sessions covering frontend development, backend concepts, teamwork, and practical implementation.",
                "Guide students through real-world development practices to improve project completion quality and technical confidence.",
            ],
        },
        {
            "title": "Python Development Intern",
            "company": "Flipkart Launchpad Student Internship Programme",
            "date": "Aug 2025 - Jan 2026",
            "location": "Remote",
            "bullets": [
                "Completed a 6-month Python Development internship organized by Corvyx and Flipkart.",
                "Built and practiced Python software development tasks in an e-commerce-oriented learning environment.",
                "Applied debugging, programming fundamentals, and software development practices across internship assignments.",
            ],
        },
    ],
    "projects": [
        {
            "name": "Employee Leave Management System (ELMS)",
            "tech": "Python, Flask, REST APIs, HTML, CSS, JavaScript",
            "link": "github.com/rkotesh/elms | elms-3.onrender.com",
            "bullets": [
                "Built a full-stack leave management system with employee and manager workflows, authentication, and responsive screens.",
                "Implemented REST API patterns to digitize leave requests, approvals, and status tracking.",
            ],
        },
        {
            "name": "College ERP Portal",
            "tech": "Django, Python, HTML, CSS, JavaScript",
            "link": "github.com/rkotesh/ciet_erp",
            "bullets": [
                "Developed an academic ERP portal for student attendance, records, assignments, and administrative workflows.",
                "Created structured dashboard screens and data-management flows for college operations.",
            ],
        },
        {
            "name": "Rolla AI Digital Agency Platform",
            "tech": "MERN Stack, Django, Next.js, Tailwind CSS",
            "link": "github.com/rkotesh/rolla-ai | rolla-ai.vercel.app",
            "bullets": [
                "Built a responsive digital agency and SaaS development platform with modern UI, SEO structure, and deployment-ready pages.",
                "Designed reusable web sections for service positioning, client acquisition, and software product presentation.",
            ],
        },
        {
            "name": "Hospital Chatbot",
            "tech": "Python, AI, Streamlit, Chatbot",
            "link": "github.com/rkotesh/hospital_chatbot | hospitalchatbot04.streamlit.app",
            "bullets": [
                "Created an AI chatbot prototype to answer patient queries and support appointment-related interactions.",
                "Built a Streamlit interface for fast browser-based use and live demonstration.",
            ],
        },
        {
            "name": "QR Code Generator",
            "tech": "Python, Streamlit",
            "link": "github.com/rkotesh/QR-Generator | qr-generator04.streamlit.app",
            "bullets": [
                "Developed a web utility that generates downloadable QR codes for URLs, text, and common input formats.",
            ],
        },
        {
            "name": "Personal Developer Portfolio",
            "tech": "HTML, CSS, JavaScript, Vercel",
            "link": "github.com/rkotesh/kotesh-portfolio | kotesh-portfolio-nine.vercel.app",
            "bullets": [
                "Built a responsive portfolio with project sections, certifications, contact workflow, and an interactive assistant.",
            ],
        },
    ],
    "education": [
        "B.Tech in Artificial Intelligence and Machine Learning, Chalapathi Institute of Engineering and Technology, 2023 - 2027 | CGPA: 8.03",
        "Intermediate, Sri Saraswathi Jr College, 2021 - 2023 | 90.07%",
        "SSC, Govt Z Z P H School, 2020 - 2021 | 97.00%",
    ],
    "certifications": [
        "Flipkart Launchpad Internship Programme - Python Development, Jan 2026",
        "Freedom With AI Master Class - AI Tools and Prompt Engineering Workflows, Jun 2025",
        "MERN Stack Training - Chalapathi Institute of Engineering and Technology, 2025",
    ],
    "achievements": [
        "Published 12 public GitHub repositories across Python, JavaScript, HTML/CSS, and TypeScript.",
        "Built and deployed live applications using Vercel, Render, and Streamlit.",
        "Maintained professional LinkedIn presence with 735 followers and 500+ connections.",
    ],
}


def set_run_font(run, name="Arial", size=10, bold=False, color="000000"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def style_doc(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(9.4)
    normal.paragraph_format.space_after = Pt(2)
    normal.paragraph_format.line_spacing = 1.04

    for style_name, size in [("Heading 1", 11), ("Heading 2", 10)]:
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(6)
        style.paragraph_format.space_after = Pt(2)


def add_bottom_border(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "BFBFBF")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_section(doc, title):
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    run = p.add_run(title.upper())
    set_run_font(run, size=10.5, bold=True)
    add_bottom_border(p)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    p.paragraph_format.space_after = Pt(1.4)
    run = p.add_run(text)
    set_run_font(run, size=9.2)


def add_docx():
    doc = Document()
    style_doc(doc)

    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name.add_run(RESUME["name"].upper())
    set_run_font(run, size=17, bold=True)
    name.paragraph_format.space_after = Pt(1)

    headline = doc.add_paragraph()
    headline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = headline.add_run(RESUME["headline"])
    set_run_font(run, size=10.5, bold=True)
    headline.paragraph_format.space_after = Pt(1)

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = contact.add_run(RESUME["contact"])
    set_run_font(run, size=8.4)
    contact.paragraph_format.space_after = Pt(5)

    add_section(doc, "Professional Summary")
    p = doc.add_paragraph()
    p.add_run(RESUME["summary"])
    set_run_font(p.runs[0], size=9.2)

    add_section(doc, "Technical Skills")
    for item in RESUME["skills"]:
        add_bullet(doc, item)

    add_section(doc, "Professional Experience")
    for role in RESUME["experience"]:
        p = doc.add_paragraph()
        set_run_font(p.add_run(role["title"]), size=9.6, bold=True)
        set_run_font(p.add_run(f" | {role['company']} | {role['date']} | {role['location']}"), size=9.3)
        p.paragraph_format.space_after = Pt(1)
        for bullet in role["bullets"]:
            add_bullet(doc, bullet)

    add_section(doc, "Projects")
    for project in RESUME["projects"]:
        p = doc.add_paragraph()
        set_run_font(p.add_run(project["name"]), size=9.6, bold=True)
        set_run_font(p.add_run(f" | {project['tech']}"), size=9.2)
        p.paragraph_format.space_after = Pt(0)
        link = doc.add_paragraph()
        set_run_font(link.add_run(project["link"]), size=8.4, color="404040")
        link.paragraph_format.space_after = Pt(0)
        for bullet in project["bullets"]:
            add_bullet(doc, bullet)

    add_section(doc, "Education")
    for item in RESUME["education"]:
        add_bullet(doc, item)

    add_section(doc, "Certifications")
    for item in RESUME["certifications"]:
        add_bullet(doc, item)

    add_section(doc, "Achievements")
    for item in RESUME["achievements"]:
        add_bullet(doc, item)

    doc.save(DOCX_PATH)


def pdf_paragraph(text, style):
    return Paragraph(text.replace("&", "&amp;"), style)


def add_pdf_section(story, styles, title):
    story.append(Spacer(1, 5))
    story.append(pdf_paragraph(title.upper(), styles["section"]))


def add_pdf_bullet(story, styles, text):
    story.append(pdf_paragraph(f"- {text}", styles["bullet"]))


def add_pdf():
    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=letter,
        rightMargin=0.62 * inch,
        leftMargin=0.62 * inch,
        topMargin=0.48 * inch,
        bottomMargin=0.48 * inch,
    )

    base = getSampleStyleSheet()
    styles = {
        "name": ParagraphStyle(
            "Name",
            parent=base["Normal"],
            fontName="Helvetica-Bold",
            fontSize=16,
            leading=18,
            alignment=1,
            spaceAfter=1,
        ),
        "headline": ParagraphStyle(
            "Headline",
            parent=base["Normal"],
            fontName="Helvetica-Bold",
            fontSize=9.6,
            leading=11,
            alignment=1,
            spaceAfter=1,
        ),
        "contact": ParagraphStyle(
            "Contact",
            parent=base["Normal"],
            fontName="Helvetica",
            fontSize=7.7,
            leading=9,
            alignment=1,
            spaceAfter=5,
        ),
        "body": ParagraphStyle("Body", parent=base["Normal"], fontName="Helvetica", fontSize=8.5, leading=10.1, spaceAfter=2),
        "section": ParagraphStyle(
            "Section",
            parent=base["Normal"],
            fontName="Helvetica-Bold",
            fontSize=9.5,
            leading=11,
            borderWidth=0.35,
            borderColor=colors.HexColor("#BFBFBF"),
            borderPadding=1.2,
            spaceBefore=3,
            spaceAfter=3,
        ),
        "role": ParagraphStyle("Role", parent=base["Normal"], fontName="Helvetica-Bold", fontSize=8.8, leading=10.2, spaceAfter=1),
        "link": ParagraphStyle("Link", parent=base["Normal"], fontName="Helvetica", fontSize=7.6, leading=8.5, textColor=colors.HexColor("#404040"), spaceAfter=0),
        "bullet": ParagraphStyle("Bullet", parent=base["Normal"], fontName="Helvetica", fontSize=8.25, leading=9.6, leftIndent=10, firstLineIndent=-6, spaceAfter=0.8),
    }

    story = [
        pdf_paragraph(RESUME["name"].upper(), styles["name"]),
        pdf_paragraph(RESUME["headline"], styles["headline"]),
        pdf_paragraph(RESUME["contact"], styles["contact"]),
    ]

    add_pdf_section(story, styles, "Professional Summary")
    story.append(pdf_paragraph(RESUME["summary"], styles["body"]))

    add_pdf_section(story, styles, "Technical Skills")
    for item in RESUME["skills"]:
        add_pdf_bullet(story, styles, item)

    add_pdf_section(story, styles, "Professional Experience")
    for role in RESUME["experience"]:
        story.append(pdf_paragraph(f"{role['title']} | {role['company']} | {role['date']} | {role['location']}", styles["role"]))
        for bullet in role["bullets"]:
            add_pdf_bullet(story, styles, bullet)

    add_pdf_section(story, styles, "Projects")
    for project in RESUME["projects"]:
        story.append(pdf_paragraph(f"{project['name']} | {project['tech']}", styles["role"]))
        story.append(pdf_paragraph(project["link"], styles["link"]))
        for bullet in project["bullets"]:
            add_pdf_bullet(story, styles, bullet)

    add_pdf_section(story, styles, "Education")
    for item in RESUME["education"]:
        add_pdf_bullet(story, styles, item)

    add_pdf_section(story, styles, "Certifications")
    for item in RESUME["certifications"]:
        add_pdf_bullet(story, styles, item)

    add_pdf_section(story, styles, "Achievements")
    for item in RESUME["achievements"]:
        add_pdf_bullet(story, styles, item)

    doc.build(story)


if __name__ == "__main__":
    add_docx()
    add_pdf()
    print(f"Wrote {DOCX_PATH}")
    print(f"Wrote {PDF_PATH}")
